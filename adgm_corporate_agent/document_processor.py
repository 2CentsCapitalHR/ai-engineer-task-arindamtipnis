import docx
import os

def add_comment(paragraph, text, author="Corporate Agent"):
    """Adds a comment to a paragraph in the document."""
    comment = paragraph.add_comment(text, author=author)
    return comment

def process_document(file_path, analysis_result):
    """
    Reads a .docx file, adds comments based on the analysis,
    and saves a new version.
    """
    try:
        doc = docx.Document(file_path)

        # This is a placeholder for the logic to find the right place to add comments.
        # For now, we'll add comments to paragraphs that contain specific keywords from the analysis.
        if analysis_result and analysis_result.get("issues_found"):
            for issue in analysis_result["issues_found"]:
                section_keyword = issue.get("section")
                comment_text = f"Issue: {issue.get('issue')}\nSuggestion: {issue.get('suggestion')}"

                for para in doc.paragraphs:
                    if section_keyword and section_keyword.lower() in para.text.lower():
                        add_comment(para, comment_text)
                        # Break after finding the first matching paragraph to avoid duplicate comments
                        break

        # Create outputs directory if it doesn't exist
        outputs_dir = "outputs"
        os.makedirs(outputs_dir, exist_ok=True)

        # Save the reviewed document
        base_name = os.path.basename(file_path)
        new_file_name = f"reviewed_{base_name}"
        output_path = os.path.join(outputs_dir, new_file_name)
        doc.save(output_path)

        return output_path

    except Exception as e:
        print(f"Error processing document {file_path}: {e}")
        return None

def identify_document_type(file_path):
    """
    Identifies the type of a document based on keywords in its content.
    """
    try:
        doc = docx.Document(file_path)
        # Read the first few paragraphs for efficiency, convert to lower case for matching
        text_content = "\n".join([p.text for p in doc.paragraphs[:20]]).lower()

        # Keyword mapping for document types
        # The keys are lowercase for case-insensitive matching
        doc_type_map = {
            "articles of association": "Articles of Association",
            "memorandum of association": "Memorandum of Association",
            "board resolution": "Board Resolution",
            "shareholder resolution": "Shareholder Resolution",
            "incorporation application": "Incorporation Application",
            "ubo declaration": "UBO Declaration",
            "register of members": "Register of Members and Directors",
            "employment contract": "Employment Contract",
            "change of registered address": "Change of Registered Address Notice",
        }

        for keyword, doc_type in doc_type_map.items():
            if keyword in text_content:
                print(f"Identified document '{os.path.basename(file_path)}' as: {doc_type}")
                return doc_type
        
        print(f"Could not identify document type for '{os.path.basename(file_path)}'.")
        return "Unknown Document Type"

    except Exception as e:
        print(f"Error identifying document type for {file_path}: {e}")
        return "Unknown Document Type"
